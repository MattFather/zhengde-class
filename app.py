import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io
import datetime
import subprocess
import tempfile
import os
import streamlit.components.v1 as components
import base64
import requests 

# ================= 1. 頁面基本設定與 JS 快捷鍵 =================
st.set_page_config(
    page_title="正德調課小幫手",
    page_icon="🏫",
    layout="wide"
)

components.html(
    """
    <script>
    const doc = window.parent.document;
    doc.addEventListener('keydown', function(event) {
        if (event.key.toLowerCase() === 'c') {
            const activeElement = doc.activeElement;
            const isInput = activeElement.tagName === 'INPUT' || activeElement.tagName === 'TEXTAREA';
            if (!isInput) {
                event.preventDefault();
                event.stopPropagation();
                event.stopImmediatePropagation();
            }
        }
    }, true); 
    </script>
    """,
    height=0,
    width=0,
)

# ================= 2. 核心資料載入 =================
@st.cache_data
def load_data():
    return pd.read_csv("schedule.csv")

try:
    df = load_data()
except FileNotFoundError:
    st.error("找不到 schedule.csv 檔案。")
    st.stop()

# ================= 3. 調課核心演算法 =================
def find_all_swaps(df, my_name, target_class, my_day, my_period):
    def is_free(teacher, check_day, check_period, ignore_day, ignore_period):
        if check_day == ignore_day and check_period == ignore_period:
            return True
        return df[(df['Teacher'] == teacher) & (df['Day'] == check_day) & (df['Period'] == check_period)].empty

    options = {}
    class_x_schedule = df[df['Class'] == target_class]

    for _, row_b in class_x_schedule.iterrows():
        teacher_b = row_b['Teacher']
        day_b = row_b['Day']
        period_b = row_b['Period']
        subject_x = row_b['Subject']

        if teacher_b == my_name: continue
        if not is_free(my_name, day_b, period_b, my_day, my_period): continue

        tb_key = f"{day_b}_{period_b}"
        if tb_key not in options:
            options[tb_key] = {
                "Teacher_B": teacher_b, "Day_B": day_b, "Period_B": period_b,
                "Subject_X": subject_x, "direct": False, "chains": []
            }

        if is_free(teacher_b, my_day, my_period, day_b, period_b):
            options[tb_key]["direct"] = True
        else:
            b_conflict = df[(df['Teacher'] == teacher_b) & (df['Day'] == my_day) & (df['Period'] == my_period)]
            if not b_conflict.empty:
                class_w = b_conflict.iloc[0]['Class']
                subject_w = b_conflict.iloc[0]['Subject']

                if class_w != target_class:
                    for _, row_c in df[df['Class'] == class_w].iterrows():
                        teacher_c = row_c['Teacher']
                        day_c = row_c['Day']
                        period_c = row_c['Period']

                        if teacher_c in [my_name, teacher_b]: continue
                        
                        if not is_free(teacher_b, day_c, period_c, day_b, period_b): continue
                        if not is_free(teacher_c, my_day, my_period, day_c, period_c): continue

                        options[tb_key]["chains"].append({
                            "type": "chain", "Teacher_C": teacher_c, "Day_C": day_c, "Period_C": period_c,
                            "Class_W": class_w, "Subject_W": subject_w, "Subject_C_W": row_c['Subject'] 
                        })

                for _, row_c in df[df['Class'] == target_class].iterrows():
                    teacher_c = row_c['Teacher']
                    day_c = row_c['Day']
                    period_c = row_c['Period']

                    if teacher_c in [my_name, teacher_b]: continue
                    
                    if not is_free(teacher_b, day_c, period_c, day_b, period_b): continue
                    if not is_free(teacher_c, my_day, my_period, day_c, period_c): continue

                    options[tb_key]["chains"].append({
                        "type": "triangle", "Teacher_C": teacher_c, "Day_C": day_c, "Period_C": period_c,
                        "Class_W": target_class, "Subject_W": row_c['Subject'] 
                    })

    return {k: v for k, v in options.items() if v["direct"] or v["chains"]}

def create_schedule_grid(df, teacher_name):
    t_df = df[df['Teacher'] == teacher_name].copy()
    all_days = ['Mon', 'Tue', 'Wed', 'Thu', 'Fri']
    all_periods = list(range(1, 9)) 
    
    if t_df.empty:
        grid = pd.DataFrame("", index=all_periods, columns=all_days)
    else:
        t_df = t_df.drop_duplicates(subset=['Period', 'Day'])
        t_df['Cell'] = t_df['Class'].astype(str) + "班\n" + t_df['Subject']
        grid = t_df.pivot(index='Period', columns='Day', values='Cell')
        grid = grid.reindex(index=all_periods, columns=all_days).fillna("")
        
    grid.index = [f"第 {i} 節" for i in all_periods]
    grid.columns = ['星期一', '星期二', '星期三', '星期四', '星期五']
    return grid

def get_next_weekday(day_zh):
    day_map = {'星期一': 0, '星期二': 1, '星期三': 2, '星期四': 3, '星期五': 4}
    target_wd = day_map.get(day_zh, 0)
    today = datetime.date.today()
    days_ahead = target_wd - today.weekday()
    if days_ahead <= 0: days_ahead += 7
    return today + datetime.timedelta(days_ahead)

# ================= 4. 列印系統演算法 =================
def docx_to_pdf(docx_bytes):
    with tempfile.TemporaryDirectory() as tmpdir:
        docx_path = os.path.join(tmpdir, "temp.docx")
        pdf_path = os.path.join(tmpdir, "temp.pdf")
        with open(docx_path, "wb") as f: f.write(docx_bytes)
        try:
            subprocess.run(["libreoffice", "--headless", "--convert-to", "pdf", "--outdir", tmpdir, docx_path], check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
            if os.path.exists(pdf_path):
                with open(pdf_path, "rb") as f: return f.read()
            return None
        except Exception: return None

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    for side in ["top", "left", "bottom", "right"]:
        if side in kwargs:
            tag = 'w:{}'.format(side)
            element = tcBorders.find(qn(tag))
            if element is not None: tcBorders.remove(element)
            element = OxmlElement(tag)
            for key, val in kwargs[side].items(): element.set(qn('w:{}'.format(key)), str(val))
            tcBorders.append(element)

def set_chinese_font(doc, font_name='標楷體'):
    doc.styles['Normal'].font.name = font_name
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

# ----------------- 專屬教務處的直式公版表單 -----------------
def add_official_form(doc, sch_year, my_name, leave_type, form_rows):
    section = doc.sections[0]
    section.orient = WD_ORIENT.PORTRAIT
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(1.0)
    section.right_margin = Cm(1.0)
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)

    copies = [("第一聯", "請假人保存"), ("第二聯", "教學組保存")]

    valid_dates = [r['o_date'] for r in form_rows if pd.notnull(r['o_date'])]
    if valid_dates:
        min_date = min(valid_dates)
        max_date = max(valid_dates)
        date_str = f"{min_date.month}  月  {min_date.day}  日至  {max_date.month}  月  {max_date.day}  日"
    else:
        date_str = "    月    日至    月    日"

    display_name = my_name if my_name and str(my_name).strip() != "None" else "_____________"

    for idx, (copy_num, copy_desc) in enumerate(copies):
        p_title = doc.add_paragraph()
        p_title.paragraph_format.space_before = Pt(0)
        p_title.paragraph_format.space_after = Pt(0)
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tab_stops = p_title.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Cm(17.5), WD_TAB_ALIGNMENT.RIGHT)

        run_t1 = p_title.add_run(f"新北市立正德國民中學      {sch_year}學年度教師自行調補代課單\t{copy_num}")
        run_t1.bold = True
        run_t1.font.size = Pt(14)
        p_title.add_run(f"\n\t{copy_desc}")

        p_sub = doc.add_paragraph()
        p_sub.paragraph_format.space_before = Pt(0)
        p_sub.paragraph_format.space_after = Pt(6)  
        p_sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run_s1 = p_sub.add_run("       教師 ")
        run_s2 = p_sub.add_run(f"{display_name}")
        run_s2.underline = True
        
        if leave_type:
            run_s3 = p_sub.add_run("      ")
            run_s4 = p_sub.add_run(f"{leave_type}")
            run_s4.underline = True
            run_s5 = p_sub.add_run(f"    日期：自  {date_str}")
            for r in [run_s1, run_s2, run_s3, run_s4, run_s5]: r.font.size = Pt(12)
        else:
            run_s3 = p_sub.add_run(f"      假    日期：自  {date_str}")
            for r in [run_s1, run_s2, run_s3]: r.font.size = Pt(12)

        table = doc.add_table(rows=7, cols=4)
        table.style = 'Table Grid'
        table.autofit = False

        widths = [Cm(1.5), Cm(1.5), Cm(3.8), Cm(12.2)]
        for j, w in enumerate(widths):
            table.columns[j].width = w
            for cell in table.columns[j].cells:
                cell.width = w

        headers = ["班級", "科目", "時      間", "異            動            情            形"]
        for c_idx, h in enumerate(headers):
            cell = table.cell(0, c_idx)
            cell.text = h
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.paragraphs[0].paragraph_format.space_before = Pt(2)
            cell.paragraphs[0].paragraph_format.space_after = Pt(2)

        table.rows[0].height = Cm(0.8)
        table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

        for r_idx in range(6):
            r = r_idx + 1
            table.rows[r].height = Cm(1.3)
            table.rows[r].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

            cell_class = table.cell(r, 0)
            cell_subj = table.cell(r, 1)
            cell_time = table.cell(r, 2)
            cell_desc = table.cell(r, 3)

            if r_idx < len(form_rows):
                row_data = form_rows[r_idx]
                cell_class.text = str(row_data['class'])
                cell_subj.text = str(row_data['subject'])

                o_date = row_data['o_date']
                if pd.notnull(o_date):
                    o_w_map = {0:"一", 1:"二", 2:"三", 3:"四", 4:"五", 5:"六", 6:"日"}
                    o_m, o_d = o_date.month, o_date.day
                    o_w = o_w_map.get(o_date.weekday(), " ")
                else:
                    o_m, o_d, o_w = "  ", "  ", "  "

                cell_time.text = f"  {o_m} 月  {o_d} 日\n星期  {o_w}  第 {row_data['o_period']} 節"

                if row_data['type'] == "調課":
                    t_date = row_data['t_date']
                    if pd.notnull(t_date):
                        t_m, t_d = t_date.month, t_date.day
                        t_w = o_w_map.get(t_date.weekday(), " ")
                    else:
                        t_m, t_d, t_w = "  ", "  ", "  "
                    t_p = row_data['t_period']
                    t_teacher = row_data['t_teacher']

                    o_teacher = str(row_data.get('o_teacher', my_name)).strip()
                    cell_desc.text = f"☑ 1. {o_teacher}老師與  {t_m} 月  {t_d} 日星期  {t_w}  第 {t_p} 節   {t_teacher}   教師調課\n☐ 2. 請 __________________________________教師代課"
                else:
                    t_teacher = row_data['t_teacher']
                    cell_desc.text = f"☐ 1. 與 ___ 月 ___ 日星期 ___ 第 ___ 節 ____________ 教師調課\n☑ 2. 請           {t_teacher}           教師代課"
            else:
                cell_time.text = "___月___日\n星期___ 第___節"
                cell_desc.text = "1. 與 ___ 月 ___ 日星期 ___ 第 ___ 節 ____________ 教師調課\n2. 請 _________________________________教師代課"

            for c_idx in range(4):
                cell = table.cell(r, c_idx)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                for p in cell.paragraphs:
                    if c_idx < 3:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    else:
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    p.paragraph_format.space_after = Pt(2)
                    p.paragraph_format.space_before = Pt(2)
                    for run in p.runs:
                        run.font.size = Pt(10.0)

        if idx == 0:
            sep = doc.add_paragraph("-" * 90)
            sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sep.paragraph_format.space_before = Pt(12)
            sep.paragraph_format.space_after = Pt(12)

# ----------------- 教師與班級通知聯 (橫式) -----------------
def generate_timetable_block(container_cell, title_suffix, sch_year, sch_term, issue_unit, class_label, filtered_df, is_teacher_side=True, teacher_name=""):
    p_header = container_cell.paragraphs[0]
    p_header.paragraph_format.space_before = Pt(0)
    p_header.paragraph_format.space_after = Pt(0)
    p_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_h = p_header.add_run(f"新北市立正德國民中學 {sch_year}學年度第{sch_term}學期\n調/代 課單")
    run_h.bold = True
    run_h.font.size = Pt(14) 

    p_sub = container_cell.add_paragraph()
    p_sub.paragraph_format.space_before = Pt(0)
    p_sub.paragraph_format.space_after = Pt(0)
    tab_stops = p_sub.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Cm(13.32), WD_TAB_ALIGNMENT.RIGHT)
    left_text = f"教師：{teacher_name}" if teacher_name else ""
    run_sub = p_sub.add_run(f"{left_text}\t班級：{class_label}")
    run_sub.bold = True
    run_sub.font.size = Pt(12) 

    inner_table = container_cell.add_table(rows=9, cols=6)
    inner_table.style = 'Table Grid'
    inner_table.autofit = False 
    inner_widths = [Cm(2.22), Cm(2.22), Cm(2.22), Cm(2.22), Cm(2.22), Cm(2.22)]
    for j, width in enumerate(inner_widths):
        inner_table.columns[j].width = width
        for cell in inner_table.columns[j].cells: cell.width = width
    
    weekdays_list = ["一", "二", "三", "四", "五"]
    h_cells = inner_table.rows[0].cells
    p_tl = h_cells[0].paragraphs[0]
    p_tl.text = ""
    run_tl = p_tl.add_run("節次/星期")
    run_tl.font.size = Pt(11)
    
    for i, day in enumerate(weekdays_list):
        p_day = h_cells[i+1].paragraphs[0]
        p_day.text = ""
        run_day = p_day.add_run(day)
        run_day.font.size = Pt(12)

    periods_list = ["1", "2", "3", "4", "5", "6", "7", "8"]
    times_list = ["08:20-09:05", "09:15-10:00", "10:10-10:55", "11:05-11:50", "13:00-13:45", "13:55-14:40", "15:00-15:45", "15:55-16:40"]
    for r_idx in range(8):
        row = inner_table.rows[r_idx + 1]
        row.height = Cm(1.5)
        cell_p = row.cells[0].paragraphs[0]
        cell_p.paragraph_format.space_after = Pt(0)
        run_num = cell_p.add_run(periods_list[r_idx])
        run_num.bold = True
        run_num.font.size = Pt(12) 
        cell_p.add_run("\n")
        run_time = cell_p.add_run(times_list[r_idx])
        run_time.font.size = Pt(9) 

    day_map = {0: 1, 1: 2, 2: 3, 3: 4, 4: 5} 
    cell_records = {}
    
    for _, row_data in filtered_df.iterrows():
        if pd.notnull(row_data["日期"]) and row_data["日期"] != "":
            d_idx = day_map.get(pd.to_datetime(row_data["日期"]).weekday())
            try: p_idx = int(str(row_data["節次"]).split()[1])
            except: continue
            if d_idx and p_idx:
                k = (p_idx, d_idx)
                if k not in cell_records: cell_records[k] = []
                cell_records[k].append(row_data)
                
    for (p_idx, d_idx), records in cell_records.items():
        cell = inner_table.rows[p_idx].cells[d_idx]
        cell.text = "" 
        
        actual_classes = [r for r in records if str(r.get("調/代課")) != "空堂X"]
        x_marks = [r for r in records if str(r.get("調/代課")) == "空堂X"]
        
        is_first = True
        
        if actual_classes:
            for idx, row_data in enumerate(actual_classes):
                if is_first: 
                    p1 = cell.paragraphs[0]
                    is_first = False
                else:
                    cell.add_paragraph()
                    p1 = cell.add_paragraph()
                    
                c_name = str(row_data["班級"]).strip() if pd.notnull(row_data["班級"]) and row_data["班級"] != "" else ""
                s_name = str(row_data["科目"]).strip() if pd.notnull(row_data["科目"]) and row_data["科目"] != "" else ""
                
                p1.paragraph_format.space_after = Pt(0)
                run_date_cell = p1.add_run(pd.to_datetime(row_data["日期"]).strftime("%m/%d"))
                run_date_cell.font.size = Pt(9)
                run_date_cell.bold = True
                
                p2 = cell.add_paragraph()
                p2.paragraph_format.space_after = Pt(0)
                
                if is_teacher_side and c_name:
                    run_c = p2.add_run(f"{c_name} ")
                    run_c.font.size = Pt(9)
                    run_c.bold = True
                
                if s_name:
                    run_s = p2.add_run(s_name)
                    run_s.bold = True
                    if len(s_name) > 4: run_s.font.size = Pt(7.5) 
                    elif len(s_name) == 4: run_s.font.size = Pt(8.0) 
                    else: run_s.font.size = Pt(9.0) 
                
                p3 = cell.add_paragraph()
                p3.paragraph_format.space_after = Pt(0)
                run_teacher = p3.add_run(str(row_data["老師"]))
                run_teacher.font.size = Pt(9)
                run_teacher.bold = True
                
                p4 = cell.add_paragraph()
                p4.paragraph_format.space_after = Pt(0)
                pair_id = str(row_data.get("配對編號", "")).strip()
                
                if str(row_data["調/代課"]) == "代課":
                    run_type = p4.add_run("[代課]")
                    run_type.font.size = Pt(8)
                else:
                    if title_suffix == "存查聯" and pair_id: run_type = p4.add_run(f"[{pair_id}]")
                    else: run_type = p4.add_run("[調課]")
                    run_type.font.size = Pt(8)
                            
        if x_marks and "教師通知聯" in title_suffix:
            for idx, row_data in enumerate(x_marks):
                if is_first: 
                    p1 = cell.paragraphs[0]
                    is_first = False
                else:
                    cell.add_paragraph()
                    p1 = cell.add_paragraph()
                    
                p1.paragraph_format.space_after = Pt(0)
                run_date_cell = p1.add_run(pd.to_datetime(row_data["日期"]).strftime("%m/%d"))
                run_date_cell.font.size = Pt(9)
                run_date_cell.bold = True
                
                p2 = cell.add_paragraph()
                p2.paragraph_format.space_after = Pt(0)
                run_x = p2.add_run("✖")
                run_x.font.size = Pt(14)
                run_x.bold = True
                
                p3 = cell.add_paragraph()
                p3.paragraph_format.space_after = Pt(0)
                target_info = str(row_data.get("原資訊", "")).strip()
                run_text = p3.add_run(target_info if target_info else "(已調走)")
                run_text.font.size = Pt(8)
                run_text.bold = True

    for r in range(9):
        for c in range(6):
            curr_cell = inner_table.rows[r].cells[c]
            curr_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for para in curr_cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.paragraph_format.line_spacing = 1.0
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after = Pt(0)
            if r == 4: set_cell_border(curr_cell, bottom={"sz": 24, "val": "single", "color": "000000"})
            if r == 5: set_cell_border(curr_cell, top={"sz": 24, "val": "single", "color": "000000"})

    print_p = container_cell.paragraphs[-1] 
    print_p.text = ""
    print_p.paragraph_format.space_before = Pt(0) 
    print_p.paragraph_format.space_after = Pt(0)
    print_p.alignment = WD_ALIGN_PARAGRAPH.LEFT 
    tab_stops_print = print_p.paragraph_format.tab_stops
    tab_stops_print.add_tab_stop(Cm(13.32), WD_TAB_ALIGNMENT.RIGHT)
    run_issue = print_p.add_run(f"發放單位：{issue_unit}")
    run_issue.font.size = Pt(10)
    run_date = print_p.add_run(f"\t列印：{datetime.date.today().strftime('%Y/%m/%d')}")
    run_date.font.size = Pt(10)

    footer_p = container_cell.add_paragraph()
    footer_p.paragraph_format.space_before = Pt(0)
    footer_p.paragraph_format.space_after = Pt(0)
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_f = footer_p.add_run(f"({title_suffix})")
    run_f.bold = True
    run_f.font.size = Pt(10)

def process_swap_logic(df):
    df_result = []
    subs = df[df["調/代課"] == "代課"].copy()
    for _, r in subs.iterrows(): df_result.append(r)
        
    swaps = df[df["調/代課"] == "調課"].copy()
    pair_ids = [p for p in swaps["配對編號"].unique() if pd.notnull(p) and str(p).strip() != ""]
    
    for pid in pair_ids:
        rows = swaps[swaps["配對編號"] == pid].sort_index()
        n = len(rows)
        if n >= 2:
            orig_dates = rows["日期"].tolist()
            orig_periods = rows["節次"].tolist()
            shifted_dates = orig_dates[1:] + [orig_dates[0]]
            shifted_periods = orig_periods[1:] + [orig_periods[0]]
            
            for i in range(n):
                new_row = rows.iloc[i].copy()
                new_row["日期"] = shifted_dates[i]
                new_row["節次"] = shifted_periods[i]
                df_result.append(new_row)
                
                x_row = rows.iloc[i].copy()
                x_row["日期"] = orig_dates[i]
                x_row["節次"] = orig_periods[i]
                x_row["調/代課"] = "空堂X"
                target_date = shifted_dates[i]
                target_period = shifted_periods[i]
                try:
                    t_date_str = pd.to_datetime(target_date).strftime('%m/%d') if pd.notnull(target_date) and target_date != "" else ""
                    t_p_num = "".join(filter(str.isdigit, str(target_period)))
                    if t_date_str and t_p_num: x_row["原資訊"] = f"調 {t_date_str}[{t_p_num}]"
                    else: x_row["原資訊"] = "(已調走)"
                except: x_row["原資訊"] = "(已調走)"
                df_result.append(x_row)
        else:
            for _, r in rows.iterrows(): df_result.append(r)
    
    no_id = swaps[swaps["配對編號"].isna() | (swaps["配對編號"] == "")]
    for _, r in no_id.iterrows(): df_result.append(r)
    return pd.DataFrame(df_result)

def create_docx(sch_year, sch_term, issue_unit, leave_type, edited_df, my_name):
    doc = Document()
    set_chinese_font(doc, '標楷體')

    df_raw = edited_df[edited_df["勾選列印資料"] == True].copy()
    if df_raw.empty: return None

    form_rows = []
    
    df_raw["調/代課"] = df_raw["調/代課"].astype(str).replace(['nan', 'None', 'NaN'], '').str.strip()
    df_raw["配對編號"] = df_raw["配對編號"].astype(str).str.replace(r'\.0$', '', regex=True).replace(['nan', 'None', 'NaN'], '').str.strip()
    df_raw["老師"] = df_raw["老師"].astype(str).replace(['nan', 'None', 'NaN'], '').str.strip()
    
    df_swaps = df_raw[df_raw["調/代課"] == "調課"]
    for pid in df_swaps["配對編號"].unique():
        if not pid or pid == "": continue
        
        rows = df_swaps[df_swaps["配對編號"] == pid]
        n = len(rows)
        
        if n == 2:
            row_o = rows.iloc[0]
            row_t = rows.iloc[1]
            form_rows.append({
                "class": row_o['班級'],
                "subject": row_o['科目'],
                "o_date": pd.to_datetime(row_o['日期']),
                "o_period": "".join(filter(str.isdigit, str(row_o['節次']))),
                "type": "調課",
                "t_date": pd.to_datetime(row_t['日期']),
                "t_period": "".join(filter(str.isdigit, str(row_t['節次']))),
                "t_teacher": row_t['老師'],
                "o_teacher": row_o['老師']
            })
        elif n >= 3:
            for i in range(n):
                row_o = rows.iloc[i]
                row_t = rows.iloc[(i + 1) % n]
                form_rows.append({
                    "class": row_o['班級'],
                    "subject": row_o['科目'],
                    "o_date": pd.to_datetime(row_o['日期']),
                    "o_period": "".join(filter(str.isdigit, str(row_o['節次']))),
                    "type": "調課",
                    "t_date": pd.to_datetime(row_t['日期']),
                    "t_period": "".join(filter(str.isdigit, str(row_t['節次']))),
                    "t_teacher": row_t['老師'],
                    "o_teacher": row_o['老師']
                })
            
    df_subs = df_raw[df_raw["調/代課"] == "代課"]
    for _, row in df_subs.iterrows():
        form_rows.append({
            "class": row['班級'],
            "subject": row['科目'],
            "o_date": pd.to_datetime(row['日期']),
            "o_period": "".join(filter(str.isdigit, str(row['節次']))),
            "type": "代課",
            "t_date": None,
            "t_period": "",
            "t_teacher": row['老師'],
            "o_teacher": my_name
        })

    add_official_form(doc, sch_year, my_name, leave_type, form_rows[:6])

    new_section = doc.add_section()
    new_section.orient = WD_ORIENT.LANDSCAPE
    new_section.page_width = Cm(29.7)
    new_section.page_height = Cm(21.0)
    new_section.left_margin = Cm(0.8)
    new_section.right_margin = Cm(0.5)
    new_section.top_margin = new_section.bottom_margin = Cm(0.5)

    df_processed = process_swap_logic(df_raw)
    all_blocks = []

    teachers = sorted(list(set([t for t in df_processed["老師"] if t != ""])))
    for t in teachers:
        df_t = df_processed[df_processed["老師"] == t]
        t_classes = sorted(list(set([c for c in df_t["班級"] if c != ""])))
        all_blocks.append({"suffix": "教師通知聯", "label": ", ".join(t_classes), "df": df_t, "is_teacher": True, "teacher_name": f"{t}老師"})

    classes = sorted(list(set([c for c in df_processed["班級"] if c != ""])))
    for c in classes:
        df_c = df_processed[df_processed["班級"] == c]
        all_blocks.append({"suffix": "班級公告聯", "label": c, "df": df_c, "is_teacher": False, "teacher_name": ""})

    for i in range(0, len(all_blocks), 2):
        if i > 0: doc.add_page_break()
        table = doc.add_table(rows=1, cols=4)
        table.autofit = False
        col_widths = [Cm(13.7), Cm(0.5), Cm(0.5), Cm(13.7)]
        for j in range(4):
            table.columns[j].width = col_widths[j]
            for cell in table.columns[j].cells: cell.width = col_widths[j]
        set_cell_border(table.cell(0, 1), right={"sz": 6, "val": "dashed", "color": "808080"})

        b1 = all_blocks[i]
        generate_timetable_block(table.cell(0, 0), b1["suffix"], sch_year, sch_term, issue_unit, b1["label"], b1["df"], is_teacher_side=b1["is_teacher"], teacher_name=b1["teacher_name"])
        if i + 1 < len(all_blocks):
            b2 = all_blocks[i+1]
            generate_timetable_block(table.cell(0, 3), b2["suffix"], sch_year, sch_term, issue_unit, b2["label"], b2["df"], is_teacher_side=b2["is_teacher"], teacher_name=b2["teacher_name"])

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

# ================= 5. 系統狀態記憶與初始化 =================
state_keys = [
    "last_user_name", "uni_source_class", "uni_source_subject", "uni_source_period", "uni_source_day_en", "uni_source_day_zh", 
    "uni_last_clicked_cell", "uni_target_tb"
]
for key in state_keys:
    if key not in st.session_state: st.session_state[key] = None

if 'res_data' not in st.session_state:
    st.session_state.res_data = pd.DataFrame({
        "勾選列印資料": pd.Series(dtype='bool'), "配對編號": pd.Series(dtype='str'), "班級": pd.Series(dtype='str'),
        "日期": pd.Series(dtype='datetime64[ns]'), "節次": pd.Series(dtype='str'), "科目": pd.Series(dtype='str'),
        "老師": pd.Series(dtype='str'), "調/代課": pd.Series(dtype='str')
    })

def check_source_conflict(df, teacher, date_val, period):
    if df.empty: return False
    date_str = pd.to_datetime(date_val).strftime('%Y-%m-%d')
    df_dates = pd.to_datetime(df['日期'], errors='coerce').dt.strftime('%Y-%m-%d')
    return not df[(df['老師'] == teacher) & (df_dates == date_str) & (df['節次'] == period)].empty

def check_destination_conflict(df, teacher, date_val, period):
    if df.empty: return False
    processed_df = process_swap_logic(df)
    if processed_df.empty: return False
    date_str = pd.to_datetime(date_val).strftime('%Y-%m-%d')
    p_dates = pd.to_datetime(processed_df['日期'], errors='coerce').dt.strftime('%Y-%m-%d')
    conflict = processed_df[(processed_df['老師'] == teacher) & (p_dates == date_str) & (processed_df['節次'] == period) & (processed_df['調/代課'] != '空堂X')]
    return not conflict.empty

def style_my_grid(val):
    val_str = str(val)
    if "🔄" in val_str: 
        return "color: #ffffff; font-weight: bold; background-color: #d9534f;" 
    elif "🌟互" in val_str: 
        return "color: #0066cc; font-weight: normal; background-color: transparent;" 
    elif "🔗多" in val_str: 
        return "color: #e67e22; font-weight: normal; background-color: transparent;" 
    elif "🌟" in val_str: 
        return "color: #0066cc; font-weight: normal; background-color: transparent;" 
    elif "班" in val_str:
        return "color: #2c3e50; font-weight: bold; background-color: #e2e8f0;" 
    return ""

def style_target_grid(val):
    val_str = str(val)
    if '\u200b' in val_str: 
        return "color: #ffffff; font-weight: bold; background-color: #d9534f;" 
    elif '\u200c' in val_str:
        return "color: #ffffff; font-weight: bold; background-color: #28a745;" 
    elif "⚠️衝堂" in val_str:
        return "color: #000000; font-weight: bold; background-color: #ffc107;" 
    elif "班" in val_str:
        return "color: #2c3e50; font-weight: bold; background-color: #e2e8f0;" 
    return ""

day_en = ['Mon', 'Tue', 'Wed', 'Thu', 'Fri']
day_zh = ['星期一', '星期二', '星期三', '星期四', '星期五']
day_map_rev = dict(zip(day_zh, day_en))
day_map_en_zh = dict(zip(day_en, day_zh))

def render_jump_button():
    jump_html = """
    <!DOCTYPE html>
    <html>
    <head>
    <style>
    body { margin: 0; padding: 0; font-family: "Source Sans Pro", sans-serif; }
    .btn {
        width: 100%;
        padding: 0.8rem;
        background-color: #28a745;
        color: white;
        border: none;
        border-radius: 0.5rem;
        cursor: pointer;
        font-size: 1.1rem;
        font-weight: bold;
        box-shadow: 0 2px 4px rgba(0,0,0,0.1);
        transition: background-color 0.2s;
    }
    .btn:hover { background-color: #218838; }
    </style>
    </head>
    <body>
    <button class="btn" onclick="
        const doc = window.parent.document;
        const tabs = doc.querySelectorAll('[data-baseweb=\\'tab\\']');
        for(let i=0; i<tabs.length; i++) {
            if(tabs[i].textContent.includes('第二步')) {
                tabs[i].click();
                doc.defaultView.scrollTo({top: 0, behavior: 'smooth'});
                break;
            }
        }
    ">👉 點此跳轉至【🖨️ 第二步：列印單據與輸出】</button>
    </body>
    </html>
    """
    components.html(jump_html, height=55)

# ================= 6. UI 版面佈局 =================

st.markdown("""
<style>
/* ================= 自訂字體大小控制區 ================= */
.main-title {
    font-size: 40px !important;  
    font-weight: 900 !important;
    line-height: 1.3 !important;
    margin-bottom: 15px !important;
    color: inherit !important;
}
.sub-title {
    font-size: 28px !important;  
    font-weight: 900 !important;
    line-height: 1.3 !important;
    margin-bottom: 15px !important;
    color: inherit !important;
}
</style>
""", unsafe_allow_html=True)


st.markdown("<div class='main-title'>🏫 正德調課小幫手</div>", unsafe_allow_html=True)

# 🌟 加上溫馨提醒小提示
st.info("**溫馨提醒 :** \n課表可能因課務有調動，老師請假時，務必和任課老師確認後，再行遞送調課申請表~", icon="💡")

col_top1, col_top2, col_top3 = st.columns([1, 2, 1])
with col_top2:
    all_teachers = sorted(df['Teacher'].dropna().unique())
    my_name = st.selectbox("🙋‍♂️ 請選擇您的名字：", all_teachers, index=None, placeholder="請選擇...")

if my_name and my_name != st.session_state.last_user_name:
    for k in state_keys: 
        st.session_state[k] = None
    st.session_state.last_user_name = my_name

    API_URL = f"https://script.google.com/macros/s/AKfycbzlk8-pGvH1S83NWfQ3ThHaLNYjTksmu81-liK0MvouHhh_FV0ZpiotOMZAgKSPNk50rw/exec?name={my_name}"
    try: requests.get(API_URL, timeout=5)
    except Exception: pass 

st.markdown("---")

tab_swap, tab_print = st.tabs(["🔄 第一步：智慧調課", "🖨️ 第二步：列印單據與輸出"])

# ----------------- Tab 1: 智慧調課 -----------------
with tab_swap:
    if my_name:
        all_swaps = {}
        if st.session_state.uni_source_class:
            all_swaps = find_all_swaps(df, my_name, st.session_state.uni_source_class, st.session_state.uni_source_day_en, st.session_state.uni_source_period)
            
        col_c1, col_c2 = st.columns([1, 1.2], gap="large")
        
        with col_c1:
            st.markdown(f"<div class='sub-title'>📅 【{my_name}老師】的課表</div>", unsafe_allow_html=True)
            
            advanced_mode = st.session_state.get("advanced_toggle", False)
            substitute_mode = st.session_state.get("substitute_toggle", False)
            
            with st.container(border=True):
                if substitute_mode:
                    st.markdown("🎯 **操作步驟：** 1️⃣ 點擊您欲請假的班級。 2️⃣ 在右側選擇指定的代課老師。", unsafe_allow_html=True)
                elif advanced_mode:
                    st.markdown("""
                        🎯 **操作步驟：** 1️⃣ 點擊想調走的班級。 2️⃣ 點擊 <span style='color: #0066cc;'><b>🌟</b></span> 或 <span style='color: #e67e22;'><b>🔗老師名字</b></span> 選擇對象。<br>
                        <span style='color: #0066cc;'><b>🌟互</b></span>：兩人互調 &emsp; | &emsp; <span style='color: #e67e22;'><b>🔗多</b></span>：跨班連鎖或三角調
                    """, unsafe_allow_html=True)
                else:
                    st.markdown("🎯 **操作步驟：** 1️⃣ 點擊想調走的班級。 2️⃣ 點擊 <span style='color: #0066cc;'><b>🌟 老師名字</b></span> 進行互調。", unsafe_allow_html=True)

            uni_my_grid = create_schedule_grid(df, my_name)
            uni_display_grid = uni_my_grid.copy()
            
            if st.session_state.uni_source_class:
                r_s = st.session_state.uni_source_period - 1
                c_s = day_zh.index(st.session_state.uni_source_day_zh)
                orig = uni_my_grid.iloc[r_s, c_s]
                uni_display_grid.iloc[r_s, c_s] = f"🔄[欲調走]\n{orig}"
                
                for tb_key, opt in all_swaps.items():
                    r = opt['Period_B'] - 1
                    c = day_en.index(opt['Day_B'])
                    if opt['direct']:
                        if advanced_mode: uni_display_grid.iloc[r, c] = f"🌟互\n{opt['Teacher_B']}\n{opt['Subject_X']}"
                        else: uni_display_grid.iloc[r, c] = f"🌟\n{opt['Teacher_B']}\n{opt['Subject_X']}"
                    elif advanced_mode: 
                        uni_display_grid.iloc[r, c] = f"🔗多\n{opt['Teacher_B']}\n{opt['Subject_X']}"
            
            try: styled_uni_grid = uni_display_grid.style.map(style_my_grid)
            except AttributeError: styled_uni_grid = uni_display_grid.style.applymap(style_my_grid)

            event_uni = st.dataframe(styled_uni_grid, use_container_width=True, height=380, on_select="rerun", selection_mode="single-cell", key="uni_schedule_grid")

            st.markdown("<br>", unsafe_allow_html=True)
            col_t1, col_t2 = st.columns(2)
            with col_t1: st.toggle("🚀 解鎖進階多角調", key="advanced_toggle")
            with col_t2: st.toggle("🆘 尋找代課老師", key="substitute_toggle")

            selection_uni = event_uni.selection.cells
            if selection_uni:
                cell = selection_uni[0]
                try:
                    r_val, c_val = (cell[0], cell[1]) if isinstance(cell, (tuple, list)) else (cell.get('row', 0) if isinstance(cell, dict) else getattr(cell, 'row', 0), cell.get('column', 0) if isinstance(cell, dict) else getattr(cell, 'column', 0))
                    c_idx = day_zh.index(str(c_val)) if str(c_val) in day_zh else int(c_val)
                    t_day_zh = day_zh[c_idx]
                    t_day_en = day_map_rev.get(t_day_zh, "Mon")
                    t_period = int(str(r_val).replace("第 ", "").replace(" 節", "")) if "第" in str(r_val) else int(r_val) + 1
                    clicked_id = f"{t_day_en}_{t_period}"

                    if st.session_state.uni_last_clicked_cell != clicked_id:
                        r_idx = t_period - 1
                        orig_content = uni_my_grid.iloc[r_idx, c_idx]
                        
                        if orig_content != "":
                            match_data = df[(df['Teacher'] == my_name) & (df['Day'] == t_day_en) & (df['Period'] == t_period)]
                            if not match_data.empty:
                                st.session_state.uni_source_class = match_data.iloc[0]['Class']
                                st.session_state.uni_source_subject = match_data.iloc[0]['Subject']
                                st.session_state.uni_source_period = t_period
                                st.session_state.uni_source_day_en = t_day_en
                                st.session_state.uni_source_day_zh = t_day_zh
                                st.session_state.uni_target_tb = None
                                st.session_state.uni_last_clicked_cell = clicked_id
                                st.rerun()
                        else:
                            if st.session_state.uni_source_class and ("🌟" in str(uni_display_grid.iloc[r_idx, c_idx]) or "🔗" in str(uni_display_grid.iloc[r_idx, c_idx])):
                                st.session_state.uni_target_tb = clicked_id
                                st.session_state.uni_last_clicked_cell = clicked_id
                                st.rerun()
                except Exception: pass
            else: st.session_state.uni_last_clicked_cell = None

        with col_c2:
            # === 代課邏輯 ===
            if substitute_mode and st.session_state.uni_source_class:
                st.markdown(f"<div class='sub-title'>🆘 安排代課老師</div>", unsafe_allow_html=True)
                all_other_teachers = [t for t in all_teachers if t != my_name]
                
                col_sub1, col_sub2 = st.columns(2)
                with col_sub1:
                    sub_sel = st.selectbox("🧑‍🏫 下拉選擇校內老師：", all_other_teachers, index=None, placeholder="下拉尋找或搜尋...")
                with col_sub2:
                    sub_txt = st.text_input("✏️ 空白表格：", placeholder="無課務師長或校外老師")
                
                sub_teacher = sub_txt.strip() if sub_txt.strip() else sub_sel
                
                if sub_teacher:
                    st.markdown(f"<div class='sub-title'>👀 {sub_teacher}老師的課表變化</div>", unsafe_allow_html=True)
                    grid_sub = create_schedule_grid(df, sub_teacher)
                    r_s = st.session_state.uni_source_period - 1
                    c_s = day_en.index(st.session_state.uni_source_day_en)
                    
                    existing_class = grid_sub.iloc[r_s, c_s]
                    if existing_class != "":
                        st.warning(f"⚠️ 注意：{sub_teacher}老師在這個時段已經有課，可能會分身乏術喔！")
                        grid_sub.iloc[r_s, c_s] = f"⚠️衝堂\n{existing_class}"
                    else:
                        grid_sub.iloc[r_s, c_s] = f"{st.session_state.uni_source_class}班\n{st.session_state.uni_source_subject}\n[代課]\u200c"
                        
                    try: st.dataframe(grid_sub.style.map(style_target_grid), use_container_width=True, height=380)
                    except AttributeError: st.dataframe(grid_sub.style.applymap(style_target_grid), use_container_width=True, height=380)
                    
                    with st.container(border=True):
                        st.markdown("<div style='font-size: 18px; font-weight: bold; margin-bottom: 10px;'>📥 確認無誤，加入列印清單</div>", unsafe_allow_html=True)
                        col_d1, col_btn = st.columns([2, 1.5])
                        with col_d1: 
                            date_mine = st.date_input(f"這節課的日期 ({st.session_state.uni_source_day_zh})", value=get_next_weekday(st.session_state.uni_source_day_zh), key="uni_d1_sub")
                        
                        with col_btn:
                            st.markdown("<br>", unsafe_allow_html=True)
                            if st.button("➕ 一鍵加入", type="primary", use_container_width=True, key="uni_btn_sub"):
                                p_m_str = f"第 {st.session_state.uni_source_period} 節"
                                conflict = False
                                
                                if check_source_conflict(st.session_state.res_data, my_name, date_mine, p_m_str):
                                    st.error(f"⚠️ 衝堂警告：您 在 {date_mine} {p_m_str} 已有安排！")
                                    conflict = True
                                elif check_source_conflict(st.session_state.res_data, sub_teacher, date_mine, p_m_str) or \
                                     check_destination_conflict(st.session_state.res_data, sub_teacher, date_mine, p_m_str):
                                    st.error(f"⚠️ 衝堂警告：{sub_teacher}老師 在 {date_mine} {p_m_str} 已有安排代/調課了！")
                                    conflict = True
                                    
                                if not conflict:
                                    new_row = pd.DataFrame([{
                                        "勾選列印資料": True, 
                                        "配對編號": "", 
                                        "班級": str(st.session_state.uni_source_class), 
                                        "日期": pd.to_datetime(date_mine), 
                                        "節次": str(p_m_str), 
                                        "科目": str(st.session_state.uni_source_subject).strip(), 
                                        "老師": str(sub_teacher).strip(), 
                                        "調/代課": "代課"
                                    }])
                                    st.session_state.res_data = pd.concat([st.session_state.res_data, new_row], ignore_index=True)
                                    st.success("✅ 代課方案已加入！")
                                    render_jump_button()

            # === 調課邏輯 ===                        
            elif st.session_state.uni_target_tb and st.session_state.uni_target_tb in all_swaps:
                opt = all_swaps[st.session_state.uni_target_tb]
                
                if opt['direct']:
                    st.markdown(f"<div class='sub-title'>👀 {opt['Teacher_B']}老師的課表變化</div>", unsafe_allow_html=True)
                    grid_b = create_schedule_grid(df, opt['Teacher_B'])
                    grid_b.iloc[opt['Period_B'] - 1, day_en.index(opt['Day_B'])] = "" 
                    grid_b.iloc[st.session_state.uni_source_period - 1, day_en.index(st.session_state.uni_source_day_en)] = f"{st.session_state.uni_source_class}班\n{opt['Subject_X']}\u200b"
                    
                    try: st.dataframe(grid_b.style.map(style_target_grid), use_container_width=True, height=380)
                    except AttributeError: st.dataframe(grid_b.style.applymap(style_target_grid), use_container_width=True, height=380)
                    
                    with st.container(border=True):
                        st.markdown("<div style='font-size: 18px; font-weight: bold; margin-bottom: 10px;'>📥 確認無誤，加入列印清單</div>", unsafe_allow_html=True)
                        col_d1, col_d2, col_btn = st.columns([2, 2, 1.5])
                        day_b_zh = [k for k, v in day_map_rev.items() if v == opt['Day_B']][0]
                        with col_d1: date_mine = st.date_input(f"您的原上課日 ({st.session_state.uni_source_day_zh})", value=get_next_weekday(st.session_state.uni_source_day_zh), key="uni_d1")
                        with col_d2: date_target = st.date_input(f"對方原上課日 ({day_b_zh})", value=get_next_weekday(day_b_zh), key="uni_d2")
                        
                        with col_btn:
                            st.markdown("<br>", unsafe_allow_html=True)
                            if st.button("➕ 一鍵加入", type="primary", use_container_width=True, key="uni_btn_direct"):
                                p_m_str = f"第 {st.session_state.uni_source_period} 節"
                                p_t_str = f"第 {opt['Period_B']} 節"
                                
                                conflict = False
                                if check_source_conflict(st.session_state.res_data, my_name, date_mine, p_m_str):
                                    st.error(f"⚠️ 衝堂：您 在 {date_mine} {p_m_str} 已有課！"); conflict = True
                                elif check_source_conflict(st.session_state.res_data, opt['Teacher_B'], date_target, p_t_str):
                                    st.error(f"⚠️ 衝堂：{opt['Teacher_B']}老師 在 {date_target} {p_t_str} 已有課！"); conflict = True
                                elif check_destination_conflict(st.session_state.res_data, my_name, date_target, p_t_str):
                                    st.error(f"⚠️ 衝堂：您 在 {date_target} {p_t_str} 已有課！"); conflict = True
                                elif check_destination_conflict(st.session_state.res_data, opt['Teacher_B'], date_mine, p_m_str):
                                    st.error(f"⚠️ 衝堂：{opt['Teacher_B']}老師 在 {date_mine} {p_m_str} 已有課！"); conflict = True
                                
                                if not conflict:
                                    current_ids = pd.to_numeric(st.session_state.res_data["配對編號"], errors='coerce').dropna()
                                    next_id = str(int(current_ids.max() + 1)) if not current_ids.empty else "1"
                                    
                                    new_rows = pd.DataFrame([
                                        {"勾選列印資料": True, "配對編號": str(next_id), "班級": str(st.session_state.uni_source_class), "日期": pd.to_datetime(date_mine), "節次": str(p_m_str), "科目": str(st.session_state.uni_source_subject).strip(), "老師": str(my_name).strip(), "調/代課": "調課"},
                                        {"勾選列印資料": True, "配對編號": str(next_id), "班級": str(st.session_state.uni_source_class), "日期": pd.to_datetime(date_target), "節次": str(p_t_str), "科目": str(opt['Subject_X']).strip(), "老師": str(opt['Teacher_B']).strip(), "調/代課": "調課"}
                                    ])
                                    st.session_state.res_data = pd.concat([st.session_state.res_data, new_rows], ignore_index=True)
                                    st.success("✅ 方案已加入！")
                                    render_jump_button()
                            
                elif advanced_mode:
                    st.markdown("<div class='sub-title' style='font-size: 22px !important;'>💡 請選擇協助的［橋樑］老師</div>", unsafe_allow_html=True)
                    
                    sorted_chains = sorted(
                        opt['chains'],
                        key=lambda c: (
                            0 if c['type'] == 'chain' else 1,  
                            c['Teacher_C'],                    
                            day_en.index(c['Day_C']),          
                            c['Period_C']                      
                        )
                    )
                    
                    # 🌟 移除選單標籤的科目名稱，保持版面清爽
                    bridge_options = {}
                    for c in sorted_chains:
                        t_type = "跨班連鎖" if c['type'] == 'chain' else "三角調"
                        d_zh = day_map_en_zh.get(c['Day_C'], "")
                        label = f"[{t_type}] {c['Teacher_C']}老師 ({d_zh}第{c['Period_C']}節)"
                        bridge_options[label] = c
                        
                    selected_bridge = st.selectbox("橋樑老師選項", list(bridge_options.keys()), label_visibility="collapsed")
                    c_data = bridge_options[selected_bridge]
                    
                    grid_b = create_schedule_grid(df, opt['Teacher_B'])
                    grid_c = create_schedule_grid(df, c_data['Teacher_C'])
                    
                    if c_data['type'] == 'chain':
                        grid_b.iloc[opt['Period_B'] - 1, day_en.index(opt['Day_B'])] = "" 
                        grid_b.iloc[st.session_state.uni_source_period - 1, day_en.index(st.session_state.uni_source_day_en)] = f"{st.session_state.uni_source_class}班\n{opt['Subject_X']}\u200b"
                        grid_b.iloc[c_data['Period_C'] - 1, day_en.index(c_data['Day_C'])] = f"{c_data['Class_W']}班\n{c_data['Subject_W']}\u200b"
                        grid_c.iloc[c_data['Period_C'] - 1, day_en.index(c_data['Day_C'])] = "" 
                        grid_c.iloc[st.session_state.uni_source_period - 1, day_en.index(st.session_state.uni_source_day_en)] = f"{c_data['Class_W']}班\n{c_data['Subject_C_W']}\u200b"
                    else:
                        grid_b.iloc[opt['Period_B'] - 1, day_en.index(opt['Day_B'])] = "" 
                        grid_b.iloc[c_data['Period_C'] - 1, day_en.index(c_data['Day_C'])] = f"{st.session_state.uni_source_class}班\n{opt['Subject_X']}\u200b"
                        grid_c.iloc[c_data['Period_C'] - 1, day_en.index(c_data['Day_C'])] = "" 
                        grid_c.iloc[st.session_state.uni_source_period - 1, day_en.index(st.session_state.uni_source_day_en)] = f"{st.session_state.uni_source_class}班\n{c_data['Subject_W']}\u200b"

                    st.markdown(f"<div class='sub-title'>👀 {opt['Teacher_B']}老師的課表變化</div>", unsafe_allow_html=True)
                    try: st.dataframe(grid_b.style.map(style_target_grid), use_container_width=True, height=380)
                    except AttributeError: st.dataframe(grid_b.style.applymap(style_target_grid), use_container_width=True, height=380)

                    st.markdown(f"<div class='sub-title'>👀 {c_data['Teacher_C']}老師的課表變化</div>", unsafe_allow_html=True)
                    try: st.dataframe(grid_c.style.map(style_target_grid), use_container_width=True, height=380)
                    except AttributeError: st.dataframe(grid_c.style.applymap(style_target_grid), use_container_width=True, height=380)

                    with st.container(border=True):
                        st.markdown("<div style='font-size: 18px; font-weight: bold; margin-bottom: 10px;'>📥 確認無誤，加入列印清單</div>", unsafe_allow_html=True)
                        day_b_zh = [k for k, v in day_map_rev.items() if v == opt['Day_B']][0]
                        day_c_zh = [k for k, v in day_map_rev.items() if v == c_data['Day_C']][0]
                        
                        col_date1, col_date2, col_date3 = st.columns(3)
                        with col_date1: date_mine = st.date_input(f"您的原上課日 ({st.session_state.uni_source_day_zh})", value=get_next_weekday(st.session_state.uni_source_day_zh), key="uni_d1_adv")
                        with col_date2: date_b = st.date_input(f"{opt['Teacher_B']} 原上課日 ({day_b_zh})", value=get_next_weekday(day_b_zh), key="uni_d2_adv")
                        with col_date3: date_c = st.date_input(f"{c_data['Teacher_C']} 原上課日 ({day_c_zh})", value=get_next_weekday(day_c_zh), key="uni_d3_adv")

                        if st.button("➕ 一鍵加入", type="primary", use_container_width=True, key="uni_btn_adv"):
                            p_mine_str = f"第 {st.session_state.uni_source_period} 節"
                            p_b_str = f"第 {opt['Period_B']} 節"
                            p_c_str = f"第 {c_data['Period_C']} 節"
                            
                            conflict = False
                            if check_source_conflict(st.session_state.res_data, my_name, date_mine, p_mine_str):
                                st.error(f"⚠️ 衝堂警告：您 在 {date_mine} {p_mine_str} 已有課！"); conflict = True
                            elif check_source_conflict(st.session_state.res_data, opt['Teacher_B'], date_b, p_b_str):
                                st.error(f"⚠️ 衝堂警告：{opt['Teacher_B']}老師 在 {date_b} {p_b_str} 已有課！"); conflict = True
                            elif check_source_conflict(st.session_state.res_data, c_data['Teacher_C'], date_c, p_c_str):
                                st.error(f"⚠️ 衝堂警告：{c_data['Teacher_C']}老師 在 {date_c} {p_c_str} 已有課！"); conflict = True
                            elif check_destination_conflict(st.session_state.res_data, my_name, date_b, p_b_str):
                                st.error(f"⚠️ 目標衝堂：您 在 {date_b} {p_b_str} 已有課！"); conflict = True
                            elif check_destination_conflict(st.session_state.res_data, opt['Teacher_B'], date_c, p_c_str):
                                st.error(f"⚠️ 目標衝堂：{opt['Teacher_B']}老師 在 {date_c} {p_c_str} 已有課！"); conflict = True
                            elif check_destination_conflict(st.session_state.res_data, c_data['Teacher_C'], date_mine, p_mine_str):
                                st.error(f"⚠️ 目標衝堂：{c_data['Teacher_C']}老師 在 {date_mine} {p_mine_str} 已有課！"); conflict = True

                            if not conflict:
                                current_ids = pd.to_numeric(st.session_state.res_data["配對編號"], errors='coerce').dropna()
                                
                                if c_data['type'] == 'chain':
                                    next_id_1 = str(int(current_ids.max() + 1)) if not current_ids.empty else "1"
                                    next_id_2 = str(int(current_ids.max() + 2)) if not current_ids.empty else "2"
                                    
                                    group1 = pd.DataFrame([
                                        {"勾選列印資料": True, "配對編號": str(next_id_1), "班級": str(st.session_state.uni_source_class), "日期": pd.to_datetime(date_mine), "節次": str(p_mine_str), "科目": str(st.session_state.uni_source_subject).strip(), "老師": str(my_name).strip(), "調/代課": "調課"},
                                        {"勾選列印資料": True, "配對編號": str(next_id_1), "班級": str(st.session_state.uni_source_class), "日期": pd.to_datetime(date_b), "節次": str(p_b_str), "科目": str(opt['Subject_X']).strip(), "老師": str(opt['Teacher_B']).strip(), "調/代課": "調課"}
                                    ])
                                    group2 = pd.DataFrame([
                                        {"勾選列印資料": True, "配對編號": str(next_id_2), "班級": str(c_data['Class_W']), "日期": pd.to_datetime(date_mine), "節次": str(p_mine_str), "科目": str(c_data['Subject_W']).strip(), "老師": str(opt['Teacher_B']).strip(), "調/代課": "調課"},
                                        {"勾選列印資料": True, "配對編號": str(next_id_2), "班級": str(c_data['Class_W']), "日期": pd.to_datetime(date_c), "節次": str(p_c_str), "科目": str(c_data['Subject_C_W']).strip(), "老師": str(c_data['Teacher_C']).strip(), "調/代課": "調課"}
                                    ])
                                    st.session_state.res_data = pd.concat([st.session_state.res_data, group1, group2], ignore_index=True)
                                else:
                                    next_id = str(int(current_ids.max() + 1)) if not current_ids.empty else "1"
                                    new_rows = pd.DataFrame([
                                        {"勾選列印資料": True, "配對編號": str(next_id), "班級": str(st.session_state.uni_source_class), "日期": pd.to_datetime(date_mine), "節次": str(p_mine_str), "科目": str(st.session_state.uni_source_subject).strip(), "老師": str(my_name).strip(), "調/代課": "調課"},
                                        {"勾選列印資料": True, "配對編號": str(next_id), "班級": str(st.session_state.uni_source_class), "日期": pd.to_datetime(date_b), "節次": str(p_b_str), "科目": str(opt['Subject_X']).strip(), "老師": str(opt['Teacher_B']).strip(), "調/代課": "調課"},
                                        {"勾選列印資料": True, "配對編號": str(next_id), "班級": str(st.session_state.uni_source_class), "日期": pd.to_datetime(date_c), "節次": str(p_c_str), "科目": str(c_data['Subject_W']).strip(), "老師": str(c_data['Teacher_C']).strip(), "調/代課": "調課"}
                                    ])
                                    st.session_state.res_data = pd.concat([st.session_state.res_data, new_rows], ignore_index=True)
                                    
                                st.success("✅ 方案已加入！")
                                render_jump_button()
    else:
        st.info("👋 歡迎！請先在最上方選擇您的名字以顯示課表。")

# ----------------- Tab 2: 🖨️ 第二步：列印單據與輸出 -----------------
with tab_print:
    with st.container(border=True):
        st.markdown("<div class='sub-title'>⚙️ 單據表頭設定</div>", unsafe_allow_html=True)
        c1, c2, c3, c4 = st.columns([1, 1, 1.5, 1])
        with c1: sch_year = st.text_input("學年度", value="114")
        with c2: sch_term = st.selectbox("學期", ["一", "二"], index=1)
        
        leave_options = ["", "課務需求", "事假", "病假", "公假", "休假", "生理假", "家庭照顧假", "身心調適假", "婚假", "娩假", "喪假", "產前假", "流產假", "延長病假", "留職停薪", "陪產檢及陪產假", "骨髓或器官捐贈假", "原住民族歲時祭儀放假"]
        with c3: leave_type = st.selectbox("假別", leave_options, index=0)
        
        with c4: issue_unit = st.text_input("發放單位", value="ＯＯＯ老師")

    df_subs = df['Subject'].dropna().astype(str).str.strip().unique().tolist()
    base_subs = ["", "國文", "英文", "數學", "生物", "理化", "地科", "地理", "歷史", "公民", "體育", "健康", "視藝", "表藝", "音樂", "家政", "童軍", "輔導", "資訊", "生科", "本土語"]
    subject_list = list(dict.fromkeys(base_subs + df_subs))

    st.markdown("<div class='sub-title' style='margin-top: 20px;'>📝 待列印清單編輯區</div>", unsafe_allow_html=True)
    
    if not st.session_state.res_data.empty:
        st.session_state.res_data["日期"] = pd.to_datetime(st.session_state.res_data["日期"], errors='coerce')
        st.session_state.res_data["勾選列印資料"] = st.session_state.res_data["勾選列印資料"].astype(bool)
        for col in ["配對編號", "班級", "節次", "科目", "老師", "調/代課"]:
            st.session_state.res_data[col] = st.session_state.res_data[col].fillna("").astype(str)
            st.session_state.res_data.loc[st.session_state.res_data[col] == "nan", col] = ""

    edited_df = st.data_editor(
        st.session_state.res_data,
        key="res_data_editor",
        column_config={
            "勾選列印資料": st.column_config.CheckboxColumn("勾選"), "配對編號": st.column_config.TextColumn("配對編號"),
            "班級": st.column_config.TextColumn("班級"), "日期": st.column_config.DateColumn("日期", format="MM/DD"),
            "節次": st.column_config.SelectboxColumn("節次", options=[f"第 {i} 節" for i in range(1, 9)]),
            "科目": st.column_config.SelectboxColumn("科目", options=subject_list), "老師": st.column_config.TextColumn("老師"),
            "調/代課": st.column_config.SelectboxColumn("調/代課", options=["調課", "代課"]),
        },
        num_rows="dynamic", use_container_width=True, hide_index=True, column_order=("勾選列印資料", "配對編號", "班級", "日期", "節次", "科目", "老師", "調/代課")
    )
    
    if not edited_df.equals(st.session_state.res_data):
        st.session_state.res_data = edited_df

    c_download, _ = st.columns([2, 8])
    with c_download:
        csv_bytes = edited_df.to_csv(index=False).encode('utf-8-sig')
        st.download_button(label="💾 下載暫存檔", data=csv_bytes, file_name=f"調代課暫存_{datetime.date.today().strftime('%Y%m%d')}.csv", mime="text/csv", use_container_width=True)

    with st.expander("🛠️ 進階：本機資料恢復 (上傳暫存檔)"):
        uploaded_file = st.file_uploader("📂 若有先前下載的暫存檔 (.csv)，請在此上傳恢復：", type=["csv"])
        if uploaded_file is not None and ('last_uploaded_id' not in st.session_state or st.session_state.last_uploaded_id != uploaded_file.file_id):
            try:
                df_upload = pd.read_csv(uploaded_file, keep_default_na=False, dtype=str)
                if "日期" in df_upload.columns: df_upload["日期"] = pd.to_datetime(df_upload["日期"], errors='coerce').dt.date
                if "勾選列印資料" in df_upload.columns: df_upload["勾選列印資料"] = df_upload["勾選列印資料"].astype(str).str.lower() == 'true'
                for col in ["配對編號", "班級", "節次", "科目", "老師", "調/代課"]:
                    if col in df_upload.columns: 
                        df_upload[col] = df_upload[col].astype(str).str.replace(r'\.0$', '', regex=True).replace(['nan', 'None', 'NaN'], '').str.strip()
                st.session_state.res_data = df_upload
                st.session_state.last_uploaded_id = uploaded_file.file_id
                st.rerun() 
            except Exception as e: st.error(f"❌ 檔案讀取失敗: {e}")

    st.markdown("<div class='sub-title' style='margin-top: 20px;'>🖨️ 列印與輸出</div>", unsafe_allow_html=True)
    with st.container(border=True):
        if issue_unit.strip() == "ＯＯＯ老師": st.error("⚠️ 提醒：請在上方修改「發放單位」(預設為ＯＯＯ老師) 後，即可解鎖列印與下載功能。")
        else:
            data_docx = create_docx(sch_year, sch_term, issue_unit, leave_type, edited_df, my_name)
            if data_docx:
                col_word, col_pdf = st.columns([1, 1])
                with col_word: st.download_button("📥 下載 Word 檔 (可編輯)", data_docx, f"正德調代課單_{datetime.date.today().strftime('%Y%m%d')}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
                with col_pdf:
                    if st.button("📥 轉換並下載 PDF (手機建議)", use_container_width=True, type="primary"):
                        with st.spinner("🚀 伺服器正在努力轉換中 (約需 5~10 秒，請耐心等候)..."):
                            pdf_data = docx_to_pdf(data_docx)
                            if pdf_data:
                                st.success("✅ 轉換成功！檔案已自動下載。")
                                b64_pdf = base64.b64encode(pdf_data).decode('utf-8')
                                pdf_filename = f"正德調代課單_{datetime.date.today().strftime('%Y%m%d')}.pdf"
                                components.html(f"<script>setTimeout(function() {{ const link = window.parent.document.createElement('a'); link.href = 'data:application/octet-stream;base64,{b64_pdf}'; link.download = '{pdf_filename}'; window.parent.document.body.appendChild(link); link.click(); window.parent.document.body.removeChild(link); }}, 300);</script>", height=0, width=0)
                                st.download_button("備用：若未自動下載請點此", pdf_data, pdf_filename, mime="application/pdf", use_container_width=True)
                            else: st.error("❌ 轉換失敗，伺服器過度繁忙或缺少套件。")